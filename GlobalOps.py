import streamlit as st
import json
import pandas as pd
import re  # [NEW] 정규표현식 사용 (기계적 추출용)
from docx import Document 
from openai import OpenAI
import google.generativeai as genai
from mistralai import Mistral

# --- 페이지 설정 ---
st.set_page_config(page_title="Global Ops Sync (Hybrid)", page_icon="🛡️", layout="wide")

# ==========================================
# 🛠️ [GM 영역] 프로젝트별 금칙어 데이터베이스
# 💡 팁: Python이 인식하기 쉽게 금지어는 반드시 '작은따옴표'로 감싸주세요.
# ==========================================
PROJECT_DATABASE = {
    "🌐 공통 (Global Standard)": """
    1. Japanese: 'マジで', 'ガチ', '超', '監視'
    2. Chinese: '監視', '大力量', '牛逼'
    3. English: 'Konglish', 'Literal translations'
    """,
    
#    "⚔️ 프로젝트 A (Fantasy RPG)": """
#    - Common: 'Zombie' -> 'Infected'
#    - Japanese: '勇者' -> '継承者'
#    - English: 'Guild' -> 'Clan'
#    """,
}
# ==========================================

# --- CSS ---
st.markdown("""
<style>
    .main-header { font-size: 2.0em; font-weight: bold; color: #1E88E5; margin-bottom: 0px; }
    .correction-box { background-color: #fff9c4; padding: 15px; border-radius: 8px; border-left: 5px solid #fbc02d; margin-top: 10px; font-size: 0.95em; }
    .rule-box { background-color: #e3f2fd; padding: 10px; border-radius: 5px; font-size: 0.85em; color: #0d47a1; white-space: pre-wrap; border: 1px solid #90caf9; }
    .detected-badge { background-color: #ffcdd2; color: #c62828; padding: 2px 6px; border-radius: 4px; font-weight: bold; font-size: 0.8em; margin-right: 5px; }
</style>
""", unsafe_allow_html=True)

# --- 파일 읽기 헬퍼 ---
# --- [수정됨] 파일 읽기 헬퍼 (표 내용 포함) ---
def read_file_content(uploaded_file):
    try:
        if uploaded_file.type == "text/plain":
            return uploaded_file.getvalue().decode("utf-8")
        elif "wordprocessingml" in uploaded_file.type or uploaded_file.name.endswith('.docx'):
            doc = Document(uploaded_file)
            full_text = []
            
            # 1. 본문(Paragraphs) 읽기
            for para in doc.paragraphs:
                full_text.append(para.text)
                
            # 2. 표(Tables) 내용 읽기
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        # 셀 안에도 여러 문단이 있을 수 있음
                        cell_text = ' '.join([p.text for p in cell.paragraphs])
                        row_data.append(f"[{cell_text}]") # 셀 구분감(Bracket) 추가
                    # 행 단위로 묶어서 추가 (표 모양 흉내)
                    full_text.append(" | ".join(row_data))
            
            return '\n'.join(full_text)
        else:
            return "지원하지 않는 파일 형식입니다."
    except Exception as e:
        return f"파일 읽기 오류: {str(e)}"

# --- [NEW] 기계적 금지어 추출 및 스캔 함수 ---
def extract_quoted_words(text):
    """ 텍스트에서 '단어' 형태(따옴표 안의 내용)를 모두 추출합니다. """
    # 정규식: 작은따옴표(') 안에 있는 1글자 이상의 문자열 추출
    return re.findall(r"'([^']{1,})'", text)

def pre_scan_blacklist(content, rules_text):
    """ 
    Python이 직접 텍스트를 스캔하여 금지어가 있는지 확인합니다. 
    Returns: 발견된 금지어 리스트 (예: ['G-Star field', 'Majide'])
    """
    forbidden_words = extract_quoted_words(rules_text)
    found_violations = []
    
    # 대소문자 무시하고 검색
    content_lower = content.lower()
    
    for word in forbidden_words:
        if word.lower() in content_lower:
            found_violations.append(word)
            
    return list(set(found_violations)) # 중복 제거 후 반환

# --- 금지어 파일 파싱 ---
def parse_blacklist_file(uploaded_file):
    if uploaded_file is None: return ""
    try:
        if uploaded_file.name.endswith('.csv'): df = pd.read_csv(uploaded_file)
        else: df = pd.read_excel(uploaded_file)
        text_list = ""
        if not df.empty:
            for _, row in df.iterrows():
                lang = str(row.iloc[0]).strip() if len(row) > 0 else "Common"
                word = str(row.iloc[1]).strip() if len(row) > 1 else ""
                reason = str(row.iloc[2]).strip() if len(row) > 2 else ""
                if word and word.lower() != "nan":
                    reason_text = f"(Reason: {reason})" if reason and reason.lower() != 'nan' else ""
                    # 강제로 따옴표를 씌워서 텍스트에 추가 (Python 스캔을 위해)
                    text_list += f"- {lang}: '{word}' {reason_text}\n"
        return text_list
    except Exception as e:
        return ""

# --- AI 호출 함수 ---
def call_ai_translator(provider, api_key, system_role, user_prompt, temperature=0.1):
    try:
        if provider == "OpenAI (GPT-4o)":
            client = OpenAI(api_key=api_key)
            response = client.chat.completions.create(
                model="gpt-4o", 
                messages=[{"role": "system", "content": system_role}, {"role": "user", "content": user_prompt}],
                temperature=temperature,
                response_format={"type": "json_object"} 
            )
            return response.choices[0].message.content

        elif provider == "Google Gemini":
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel('gemini-2.5-flash') 
            response = model.generate_content(
                f"{system_role}\n\n[IMPORTANT] Output MUST be raw JSON.\n\n[User Request]\n{user_prompt}",
                generation_config=genai.types.GenerationConfig(
                    temperature=temperature, 
                    response_mime_type="application/json"
                )
            )
            return response.text

        elif provider == "Mistral AI":
            client = Mistral(api_key=api_key)
            response = client.chat.complete(
                model="mistral-small-latest",
                messages=[{"role": "system", "content": system_role}, {"role": "user", "content": user_prompt}],
                temperature=temperature,
                response_format={"type": "json_object"}
            )
            return response.choices[0].message.content
            
    except Exception as e:
        return f"Error: {str(e)}"

# --- 사이드바 ---
with st.sidebar:
    st.title("⚙️ 설정")
    provider = st.selectbox("🤖 AI 모델", ["Google Gemini", "Mistral AI", "OpenAI (GPT-4o)"])
    if provider == "Mistral AI":
        st.caption("Mistral API Key 필요")
        col_a, col_b = st.columns(2)
        with col_a: st.link_button("📘 가이드", "https://docs.google.com/presentation/d/1xTUWrusNROIonDWL5hEWpybNCqo2W8kYHr4czDPWnok/edit?slide=id.p")
        with col_b: st.link_button("🔑 키 발급", "https://console.mistral.ai/")
    api_key = st.text_input(f"{provider} API Key", type="password", placeholder="sk-...")
    st.divider()
    st.subheader("📁 프로젝트 선택")
    selected_project = st.selectbox("검수할 프로젝트", list(PROJECT_DATABASE.keys()), index=0)
    with st.expander("📜 적용된 룰 보기"):
        st.markdown(f"<div class='rule-box'>{PROJECT_DATABASE[selected_project]}</div>", unsafe_allow_html=True)
    
    # [NEW] 사용 가이드 추가
    with st.expander("❓ 사용 방법 (How to Use)"):
        st.markdown("""
        1. **파일 업로드:** 공지사항 원문(한국어)과 번역본 파일들을 **모두** 좌측에 드래그하세요.
        2. **언어 설정:** 우측에서 '기준(Master) 언어'를 선택하세요. (보통 '한국어')
        3. **금지어 추가:**
           - 프로젝트별 기본 룰은 자동 적용됩니다.
           - 추가로 검사하고 싶은 단어가 있다면 엑셀 파일을 업로드하세요.
           - **엑셀 형식:** 1열(언어), 2열(단어), 3열(이유)
        4. **분석 시작:** 버튼을 누르면 AI가 금지어 포함 여부와 톤앤매너를 검사합니다.
        5. **결과 확인:**
           - **빨간색:** 심각한 오류(금지어, 톤앤매너 붕괴) -> 전면 재작성 권장.
           - **노란색:** 부분 수정 필요.
        """)
        
# --- 메인 화면 ---
st.markdown('<div class="main-header">⚖️ 공지사항 내용 비교 도우미</div>', unsafe_allow_html=True)
col_main, col_ctrl = st.columns([1.8, 1.2])

with col_main:
    st.markdown("### 1️⃣ 분석 대상 파일")
    uploaded_files = st.file_uploader("공지사항 파일 드래그 (.txt, .docx)", type=['txt', 'docx'], accept_multiple_files=True)

with col_ctrl:
    st.markdown("### 2️⃣ 추가 옵션")
    extra_blacklist_file = st.file_uploader("➕ (선택) 임시 금지어 파일", type=['xlsx', 'csv', 'xls'])
    master_lang = st.selectbox("기준(Master) 언어", ["한국어 (Korean)", "영어 (English)", "일본어 (Japanese)", "중국어 (Chinese)"], index=0)
    deep_dive = st.toggle("🐢 심층 정밀 분석", value=True)
    st.write("") 
    analyze_btn = st.button("🚀 분석 시작", type="primary", use_container_width=True)

# --- 분석 로직 ---
if analyze_btn:
    if not uploaded_files:
        st.warning("파일을 업로드해주세요.")
    elif not api_key:
        st.error("API Key를 입력해주세요.")
    else:
        with st.spinner(f"🕵️‍♂️ AI가 [{selected_project}] 규칙에 맞춰 초정밀 분석 중입니다..."):
            
            # 1. 룰 텍스트 병합 (Python 스캔용)
            final_rules_text = PROJECT_DATABASE["🌐 공통 (Global Standard)"]
            if selected_project != "🌐 공통 (Global Standard)":
                final_rules_text += f"\n{PROJECT_DATABASE[selected_project]}"
            if extra_blacklist_file:
                extra_rules = parse_blacklist_file(extra_blacklist_file)
                final_rules_text += f"\n{extra_rules}"

            # 2. 파일 처리 및 [사전 스캔] 실행
            files_context = ""
            for up_file in uploaded_files:
                content = read_file_content(up_file)
                
                # [핵심] Python이 먼저 금지어를 찾습니다!
                detected_violations = pre_scan_blacklist(content, final_rules_text)
                violation_alert = ""
                if detected_violations:
                    # AI에게 넘겨줄 "적발 리스트" 생성
                    violation_alert = f"\n\n[🚨 SYSTEM DETECTED VIOLATIONS]\nThe system found these forbidden words in this file: {detected_violations}. You MUST report them as Critical Errors."

                files_context += f"\n\n--- FILE START ---\nFilename: {up_file.name}\nContent:\n{content}{violation_alert}\n--- FILE END ---"

            # 3. 시스템 프롬프트 (기준 언어 제외 룰 복구 & 재작성 로직 유지)
            system_role = (
                "너는 글로벌 게임 서비스의 **'가장 꼼꼼하고 지독한'** Localization QA Lead다.\n"
                "너의 임무는 텍스트를 검수하고, 상태가 심각하면 **아예 새로 작성(Total Rewrite)**하는 것이다.\n\n"
                
                "**[🚨 금지어 리스트 (Active Rules)]**\n"
                f"{final_rules_text}\n\n"
                
                "**[검수 절대 원칙]**\n"
                "1. **이중 검수 (Dual Inspection):**\n"
                "   - A. **시스템 적발:** `[🚨 SYSTEM DETECTED]`에 있는 단어는 무조건 리포트해라.\n"
                "   - B. **AI 심층 검수:** 문법, 톤앤매너, 뉘앙스 오류를 찾아라.\n"
                "2. **전면 재작성 판단 (Critical Failure Check):**\n"
                "   - 만약 텍스트가 수정 불가능할 정도로 엉망이거나(예: Little Jacob 스타일 슬랭, 기계 번역 투, 톤앤매너 완전 붕괴) 수정 범위가 50%를 넘는다면...\n"
                "   - **`critical_rewrite_needed: true`**를 설정하고 **완벽하게 새로 쓴 텍스트(full_rewrite)**를 작성해라.\n"
                "   - **[중요]** 이때 `improvements` 리스트는 **빈 배열([])**로 남겨라. (폐기할 텍스트를 굳이 고치지 마라)\n"
                "3. **평가:** 평가는 한국어로 구체적으로 작성해라.\n"
                "4. **기준 제외:** 사용자가 선택한 Master Language 파일은 **절대 분석 대상에 포함하지 마라.** (Reference ONLY)"
            )
            
            # 4. 유저 프롬프트 (줄바꿈/문단 유지 지시 추가)
            user_prompt = f"""
            [Uploaded Files]
            {files_context}
            
            **[Analysis Steps]**
            1. **Identify Master:** File for '{master_lang}'. (Reference ONLY. DO NOT Analyze.)
            2. **Identify Targets:** Analyze ALL other files (Target Languages).
            
            3. **Execute QA (For each Target file):**
               - **Step 1 (Strict Quality Check - Go/No-Go):**
                 - **CRITICAL FAILURE CONDITIONS:**
                   1. **Garbage Quality:** Slang, Broken Grammar, AI-translated feeling.
                   2. **Regional Dialects:** (e.g., 'Hello po', 'Do the needful'). CRITICAL FAILURE.
                   3. **Tone Mismatch:** Too casual or too archaic.
                   4. **EXTREME Omission:** **Only if MORE THAN 50% of the content is missing.**
                 
                 - **PASS CONDITIONS:**
                   - If the translation is generally good but misses 1~3 sentences, it is **NOT** a critical failure. -> Set `critical_rewrite_needed: false`.

                 - **Decision:**
                   - IF FAILED: 
                     - Set `critical_rewrite_needed: true`.
                     - **Write `full_rewrite` (Standard Business Professional Tone).**
                     - **[IMPORTANT] MUST preserve the exact paragraph structure and line breaks (\\n) of the Master file.**
                     - Leave `improvements` EMPTY.
                   - IF PASSED: Set `critical_rewrite_needed: false`, Proceed to Step 2.

               - **Step 2 (Detail Inspection & Omission Check):**
                 - List `improvements` for typos, wrong terms, or **Missing Sentences**.
                 - **[IMPORTANT] Handling Omissions:**
                   - If a sentence exists in Master but is MISSING in Target:
                   - `original`: "[[Master Sentence]]"
                   - `current`: "⚠️ (MISSING CONTENT)"
                   - `suggestion`: "[[Translated Sentence to add]]"
                   - `reason`: "Content omitted from Master file."
                 - **Mapping Rule:** The `"original"` field MUST ALWAYS be the Korean text from Master.
            
            **[JSON Output Format]**
            {{
                "results": [
                    {{
                        "language": "Target Language",
                        "filename": "Filename",
                        "score": "Star Rating (1~5)",
                        "tone_comparison": "평가 (한국어)",
                        "cultural_nuance": "평가 (한국어)",
                        "critical_rewrite_needed": true/false, 
                        "full_rewrite": "[[Text with strict line breaks (\\n)]]",
                        "improvements": [...]
                    }}
                ]
            }}
            """
            
            # 5. AI 호출
            temp_val = 0.0 if deep_dive else 0.1
            result_text = call_ai_translator(provider, api_key, system_role, user_prompt, temperature=temp_val)
                                
            # [결과 처리 및 UI 렌더링]
            if result_text.startswith("Error:"):
                st.error("🚨 AI 호출 중 오류가 발생했습니다.")
                st.code(result_text, language="text") 
            else:
                try:
                    data = json.loads(result_text)
                    st.success("✅ 분석 완료!")
                    
                    results = data.get("results", [])
                    if results:
                        tabs = st.tabs([f"{r['language']}" for r in results])
                        
                        # [핵심] 여기서 i (인덱스)를 활용합니다.
                        for i, tab in enumerate(tabs):
                            item = results[i]
                            with tab:
                                st.markdown(f"### 🏳️ {item['language']} ({item['filename']})")
                                st.markdown(f"**싱크로율:** {item['score']}")
                                
                                c1, c2 = st.columns(2)
                                with c1:
                                    st.markdown("**⚖️ 톤앤매너**"); st.info(item.get('tone_comparison', '-'))
                                with c2:
                                    st.markdown("**🌍 뉘앙스**"); st.info(item.get('cultural_nuance', '-'))
                                st.markdown("---")

                                # [수정됨] Key 추가하여 에러 방지
                                if item.get("critical_rewrite_needed") is True:
                                    st.error("🚨 **[심각한 품질 저하 감지]** 번역 상태가 매우 좋지 않습니다.")
                                    st.markdown("부분 수정보다는 **전면 재작성**을 권장합니다. 아래는 AI가 마스터(한국어) 파일을 기준으로 새로 작성한 초안입니다.")
                                    
                                    st.text_area(
                                        "✨ AI 추천 전면 재작성 (Full Rewrite)", 
                                        value=item.get("full_rewrite", "재작성 데이터 없음"), 
                                        height=300,
                                        key=f"rewrite_area_{i}"  # <--- [중요] 여기에 유니크 키 추가!
                                    )

                                else:
                                    # 기존 방식 (부분 수정 제안)
                                    st.markdown("**🚩 개선 제안**")
                                    imps = item.get('improvements', [])
                                    if not imps:
                                        if "WRONG" in str(item['score']):
                                            st.error("⚠️ 잘못된 파일입니다.")
                                        else:
                                            st.success("✅ 완벽합니다!")
                                    else:
                                        for imp in imps:
                                            reason_html = imp.get('reason','')
                                            if "detected" in reason_html or "금지어" in reason_html:
                                                reason_html = f'<span class="detected-badge">🚫 금지어 적발</span> {reason_html}'
                                                
                                            st.markdown(f"""
                                            <div class="correction-box">
                                                <b>🇰🇷 원문 (Master):</b> {imp.get('original','')}<br><br>
                                                <b>🚩 문제 번역:</b> {imp.get('current','')}<br><br>
                                                <b>✨ 수정 제안:</b> {imp.get('suggestion','')}<br><br>
                                                <b>💡 이유:</b> {reason_html}
                                            </div>
                                            """, unsafe_allow_html=True)
                    else:
                        st.warning("결과 없음.")

                except Exception as e:
                    st.error(f"오류: {str(e)}")



